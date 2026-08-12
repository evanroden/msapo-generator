"""
Document generator: inserts AI-extracted scope into the real RRH MSAPO template.

Template Structure (preserved as-is):
  Header:  "MSAPO AGREEMENT" / "ATTACHMENT A – SCOPE OF WORK"
  Section I:  "ADDITIONAL MSAPO DOCUMENTS" – checkbox exhibit table (untouched)
  Section II: "MSAPO SCOPE OF WORK"
              "Subcontractor shall execute the following Scope of Work
               in strict accordance with this MSAPO:"
              >>> scope content is inserted HERE, after the colon <<<

The exhibit table, header, and all formatting above the scope line are
left completely untouched.
"""

from __future__ import annotations

import re
import time
import uuid
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement

from app.config import (
    TEMPLATE_PATH,
    OUTPUT_DIR,
    FACILITY_SHORT_NAMES,
    facility_key_from_name,
)
from app.quote_analyzer import QuoteAnalysis

# The sentinel text that marks where scope content begins
SCOPE_SENTINEL = "Subcontractor shall execute the following Scope of Work in strict accordance with this MSAPO:"
_OUTPUT_MAX_AGE_SECONDS = 24 * 60 * 60
_cleanup_done = False


def _add_bullet_paragraph(doc: Document, text: str) -> "Paragraph":
    """Add a bullet-style paragraph that works even without a 'List Bullet' style."""
    try:
        para = doc.add_paragraph(style="List Bullet")
    except KeyError:
        # Template doesn't define 'List Bullet' — fake it with a dash prefix
        para = doc.add_paragraph()
        text = f"•  {text}"
    return para, text


def _add_bullet(doc, item_text: str):
    """Add a single reviewed bullet item."""
    try:
        para = doc.add_paragraph(style="List Bullet")
    except KeyError:
        para = doc.add_paragraph()
        item_text = f"•  {item_text}"
    run = para.add_run(item_text)
    return para


def _set_cell_text(cell, text: str) -> None:
    """Replace a table cell's text, preserving the first run's formatting."""
    paragraphs = cell.paragraphs
    first_para = paragraphs[0]
    # Drop any extra paragraphs in the cell
    for p in paragraphs[1:]:
        p._element.getparent().remove(p._element)
    runs = first_para.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        first_para.add_run(text)


def _clear_first_exhibit_row(doc: Document) -> None:
    """
    In the Section I "ADDITIONAL MSAPO DOCUMENTS" exhibit table, the first data
    row (Attachment A – Prime Contract) ships pre-filled with an "X" in the
    Included column and a date.  Clear both so the row matches every other row
    in the template (unchecked, blank date).  The table is otherwise untouched.
    """
    for table in doc.tables:
        if not table.rows:
            continue
        header = [c.text.strip().lower() for c in table.rows[0].cells]
        if "included" not in header or "attachment" not in header:
            continue
        if len(table.rows) < 2:
            return
        included_idx = header.index("included")
        first_row = table.rows[1]
        # Uncheck the "Included" box
        _set_cell_text(first_row.cells[included_idx], "")
        # Blank the date, copying the empty placeholder used by the other rows
        if "date" in header:
            date_idx = header.index("date")
            blank_date = ""
            if len(table.rows) > 2:
                blank_date = table.rows[2].cells[date_idx].text
            _set_cell_text(first_row.cells[date_idx], blank_date)
        return


def _find_scope_paragraph_index(doc: Document) -> int:
    """Find the paragraph index containing the scope sentinel text."""
    for i, para in enumerate(doc.paragraphs):
        if SCOPE_SENTINEL in para.text:
            return i
    raise ValueError(
        f"Could not find the scope sentinel paragraph in the template. "
        f"Expected to find: '{SCOPE_SENTINEL}'"
    )


def _insert_paragraph_after(paragraph, text: str, style=None):
    """Insert a new paragraph directly after the given paragraph element."""
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    # Manually wire up the new paragraph to the document
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, paragraph._element.getparent())
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def _append_scope_content(
    doc: Document,
    analysis: QuoteAnalysis,
    final_inclusions: list[str],
    final_exclusions: list[str],
    facility_display: str | None = None,
    facility_address_display: str | None = None,
) -> None:
    """
    Append scope content after the sentinel paragraph.

    The web review supplies the final inclusion and exclusion lists as-is.

    facility_display and facility_address_display, when given, override the
    facility values written into the document. This ensures a user's corrected
    routing choice is reflected in the attachment rather than only its filename.
    """
    # -- Facility --
    facility = facility_display or analysis.facility_name
    if facility:
        p = doc.add_paragraph()
        run = p.add_run(f"Facility: {facility}")
        run.bold = True
        run.font.size = Pt(11)
        address = (
            analysis.facility_address
            if facility_address_display is None
            else facility_address_display
        )
        doc.add_paragraph(address or "")

    # -- Vendor --
    p = doc.add_paragraph()
    run = p.add_run(f"Vendor: {analysis.vendor_name}")
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph("")  # spacer

    # -- Project Description --
    p = doc.add_paragraph()
    run = p.add_run("Project Description")
    run.bold = True
    run.font.size = Pt(11)
    doc.add_paragraph(analysis.project_description)

    doc.add_paragraph("")  # spacer

    # -- Detailed Scope --
    p = doc.add_paragraph()
    run = p.add_run("Detailed Scope of Work")
    run.bold = True
    run.font.size = Pt(11)

    for para_text in analysis.scope_of_work.split("\n"):
        stripped = para_text.strip()
        if stripped:
            doc.add_paragraph(stripped)

    doc.add_paragraph("")  # spacer

    # -- Inclusions --
    if final_inclusions:
        p = doc.add_paragraph()
        run = p.add_run("Inclusions")
        run.bold = True
        run.font.size = Pt(11)
        for item in final_inclusions:
            _add_bullet(doc, item)

    # -- Exclusions --
    if final_exclusions:
        p = doc.add_paragraph()
        run = p.add_run("Exclusions")
        run.bold = True
        run.font.size = Pt(11)
        for item in final_exclusions:
            _add_bullet(doc, item)

    doc.add_paragraph("")  # spacer

    # -- Tax Warning --
    if analysis.tax_warning:
        para = doc.add_paragraph()
        run = para.add_run(f"⚠ TAX WARNING: {analysis.tax_warning}")
        run.bold = True
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        run.font.size = Pt(11)
    elif analysis.tax_status == "included":
        para = doc.add_paragraph()
        run = para.add_run("Tax Status: Sales tax is included per the vendor quote.")
        run.font.size = Pt(10)
        run.italic = True


def _cleanup_old_outputs() -> None:
    """Best-effort cleanup of generated files older than one day.

    Run once per process: this is housekeeping rather than per-request work,
    and the sweep cost grows with the output directory.
    """
    global _cleanup_done
    if _cleanup_done:
        return
    _cleanup_done = True
    cutoff = time.time() - _OUTPUT_MAX_AGE_SECONDS
    try:
        for path in OUTPUT_DIR.iterdir():
            if path.is_file() and path.suffix.lower() in {".docx", ".pdf", ".htm"}:
                try:
                    if path.stat().st_mtime < cutoff:
                        path.unlink()
                except OSError:
                    continue
    except OSError:
        pass


def generate_docx(
    analysis: QuoteAnalysis,
    final_inclusions: list[str],
    final_exclusions: list[str],
    output_name: str | None = None,
    facility_display: str | None = None,
    facility_address_display: str | None = None,
) -> Path:
    """
    Open the MSAPO template, preserve everything at the top, and insert
    the scope content after the sentinel line.

    Args:
        analysis: The structured quote analysis from the AI.
        final_inclusions: Inclusion items retained in the web review.
        final_exclusions: Exclusion items retained in the web review.
        output_name: Optional filename stem (without extension).

    Returns the path to the generated .docx file.
    """
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(
            f"Template not found at {TEMPLATE_PATH}. "
            "Place your Master_MSAPO_Template.docx in the templates/ folder."
        )

    doc = Document(str(TEMPLATE_PATH))

    # Verify the template has the expected structure
    _find_scope_paragraph_index(doc)

    # Ensure the first exhibit row isn't pre-checked / pre-dated
    _clear_first_exhibit_row(doc)

    # Append all scope content after the existing template content
    _append_scope_content(
        doc, analysis,
        final_inclusions=final_inclusions,
        final_exclusions=final_exclusions,
        facility_display=facility_display,
        facility_address_display=facility_address_display,
    )

    # ── Build output filename ─────────────────────────────────────────
    if not output_name:
        # Pattern: "RRH {Site} {Title} MSAPO", with the site resolved from the
        # matched facility (previously any non-UMMC facility was mislabeled
        # "St. Marys" in the filename).
        fac_key = facility_key_from_name(analysis.facility_name)
        site = FACILITY_SHORT_NAMES.get(fac_key, "") if fac_key else ""
        safe_desc = re.sub(r"[^\w\s\-]", "", analysis.project_description or "SOW")[:50]
        # Don't cut mid-word ("…seals in t") — trim back to the last full word
        if len(safe_desc) == 50 and " " in safe_desc:
            safe_desc = safe_desc.rsplit(" ", 1)[0]
        parts = ["RRH", site, safe_desc.strip(), "MSAPO"]
        output_name = " ".join(p for p in parts if p)

    # Clean up filename. The random suffix prevents two concurrent sessions
    # with the same contract/site/description from overwriting one another.
    output_name = re.sub(r'[<>:"/\\|?*]', "_", output_name)
    _cleanup_old_outputs()
    unique = uuid.uuid4().hex[:10]
    docx_path = OUTPUT_DIR / f"{output_name}-{unique}.docx"
    doc.save(str(docx_path))
    return docx_path


def build_msapo_pdf(
    *,
    analysis: QuoteAnalysis,
    scope: str,
    inclusions: list[str],
    exclusions: list[str],
    facility_display: str | None = None,
    facility_address_display: str | None = None,
) -> bytes:
    """Render the official MSAPO form to PDF bytes for the PO attachment.

    Contract administration asked for the full MSAPO agreement form rather than
    the simplified Scope/Inclusions/Exclusions sheet, so this fills the real
    template and converts it, returning bytes so the caller's session-state and
    attachment plumbing is unchanged.

    ``scope`` is the operator's REVIEWED text, not ``analysis.scope_of_work``.
    The two differ whenever the scope box was edited, and rendering the raw
    analysis would silently discard those edits -- the document is the artifact
    the administrator acts on, so it must carry what the operator approved.

    Both intermediate files are removed before returning: only the PDF bytes
    travel onward, so a long-running container does not accumulate a .docx and
    .pdf per generation.
    """
    from dataclasses import replace

    from app.pdf_converter import PDFConversionError, convert_to_pdf

    reviewed = replace(analysis, scope_of_work=scope or "")
    docx_path: Path | None = None
    pdf_path: Path | None = None
    try:
        docx_path = generate_docx(
            reviewed,
            list(inclusions),
            list(exclusions),
            facility_display=facility_display,
            facility_address_display=facility_address_display,
        )
        pdf_path = convert_to_pdf(docx_path)
        payload = pdf_path.read_bytes()
    finally:
        for path in (docx_path, pdf_path):
            if path is not None:
                try:
                    path.unlink(missing_ok=True)
                except OSError:
                    pass

    if not payload.startswith(b"%PDF-"):
        raise PDFConversionError(
            "The MSAPO form did not convert to a valid PDF. The .docx renderer "
            "may be unavailable in this deployment."
        )
    return payload
