"""
Scaffold generator for Master_MSAPO_Template.docx

Run this ONCE to create a starter template with placeholder tags and
checkbox tables. Replace with your actual MSAPO template afterwards.

Usage:
    python create_template.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

TEMPLATE_DIR = Path(__file__).parent / "templates"
TEMPLATE_DIR.mkdir(exist_ok=True)
OUTPUT = TEMPLATE_DIR / "Master_MSAPO_Template.docx"


def add_checkbox_row(table, label: str, row_idx: int):
    """Add a row with a checkbox character and label."""
    row = table.rows[row_idx]
    # Checkbox cell
    row.cells[0].text = "☐"
    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Label cell
    row.cells[1].text = label


def main():
    doc = Document()

    # ── Page style ───────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Header ───────────────────────────────────────────────────────
    h = doc.add_heading("MASTER SERVICE AGREEMENT / PURCHASE ORDER", level=0)
    for run in h.runs:
        run.font.size = Pt(16)

    doc.add_paragraph("")

    # ── Top metadata table ───────────────────────────────────────────
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.style = "Light Grid Accent 1"
    meta_rows = [
        ("Date:", "{{DATE}}"),
        ("Vendor / Contractor:", "{{VENDOR}}"),
        ("Facility:", "{{FACILITY_NAME}}"),
        ("Facility Address:", "{{FACILITY_ADDRESS}}"),
        ("Project Description:", "{{PROJECT_DESCRIPTION}}"),
    ]
    for i, (label, value) in enumerate(meta_rows):
        meta_table.rows[i].cells[0].text = label
        meta_table.rows[i].cells[1].text = value
        # Bold the label
        for run in meta_table.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True

    doc.add_paragraph("")

    # ── Checkbox table (approval section) ────────────────────────────
    doc.add_heading("Approval Checklist", level=2)

    cb_items = [
        "Insurance Certificate Received",
        "W-9 Received",
        "Background Check Completed",
        "Safety Orientation Completed",
        "Scope of Work Reviewed",
        "Budget Approved",
    ]

    cb_table = doc.add_table(rows=len(cb_items), cols=2)
    cb_table.style = "Table Grid"

    # Set column widths
    for row in cb_table.rows:
        row.cells[0].width = Inches(0.5)
        row.cells[1].width = Inches(5.5)

    for i, item in enumerate(cb_items):
        add_checkbox_row(cb_table, item, i)

    doc.add_paragraph("")

    # ── Signature block ──────────────────────────────────────────────
    doc.add_heading("Signatures", level=2)

    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.style = "Table Grid"
    sig_labels = [
        ("Facilities Manager:", "Date:"),
        ("Department Head:", "Date:"),
        ("Vendor Representative:", "Date:"),
    ]
    for i, (left, right) in enumerate(sig_labels):
        sig_table.rows[i].cells[0].text = f"{left} ____________________"
        sig_table.rows[i].cells[1].text = f"{right} ____________"

    doc.add_paragraph("")
    doc.add_paragraph(
        "— Scope of Work details will be appended below this line by the "
        "MSAPO Generator —"
    ).italic = True

    # ── Save ─────────────────────────────────────────────────────────
    doc.save(str(OUTPUT))
    print(f"Template created at: {OUTPUT}")


if __name__ == "__main__":
    main()
