"""The PO attachment is the official MSAPO agreement form again.

Contract administration reversed the earlier simplification: the purchase-order
package must carry the full MSAPO form rendered to PDF, not the lightweight
Scope/Inclusions/Exclusions sheet. These tests pin the two properties that
matter and are easy to lose:

1. The document carries the operator's REVIEWED scope, not the raw analysis.
2. Rendering leaves no intermediate .docx/.pdf behind on a long-running
   container.

The full DOCX -> PDF conversion needs libreoffice-writer, which is installed by
the Dockerfile and present on CI runners but not in every sandbox, so the
end-to-end case skips itself when the renderer is genuinely unavailable rather
than reporting a failure the code did not cause.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from app.config import OUTPUT_DIR
from app.document_generator import build_msapo_pdf, generate_docx
from app.pdf_converter import PDFConversionError, convert_to_pdf
from app.quote_analyzer import QuoteAnalysis


def _analysis() -> QuoteAnalysis:
    return QuoteAnalysis(
        vendor_name="Crosby-Brownlie",
        project_description="Replace pump seals",
        scope_of_work="RAW analysis scope that the operator replaced",
        facility_name="Rochester General Hospital",
    )


def _renderer_available() -> bool:
    probe = OUTPUT_DIR / "_msapo_renderer_probe.docx"
    document = Document()
    document.add_paragraph("probe")
    document.save(str(probe))
    try:
        convert_to_pdf(probe).unlink(missing_ok=True)
        return True
    except (PDFConversionError, FileNotFoundError, OSError):
        return False
    finally:
        probe.unlink(missing_ok=True)


def test_document_uses_the_reviewed_scope_not_the_raw_analysis():
    """The operator's edits are the whole point of the review step.

    generate_docx reads analysis.scope_of_work, so the caller must inject the
    reviewed text. Rendering the raw analysis would silently discard every edit
    made in the Scope of Work box, and the document is the artifact the
    administrator acts on.
    """
    from dataclasses import replace

    reviewed = replace(
        _analysis(),
        scope_of_work="REVIEWED scope the operator approved.\nSecond line.",
    )
    path = generate_docx(
        reviewed,
        ["Labor and materials"],
        ["After-hours work"],
        facility_display="Rochester General Hospital",
    )
    try:
        text = "\n".join(p.text for p in Document(str(path)).paragraphs)
    finally:
        path.unlink(missing_ok=True)

    assert "REVIEWED scope the operator approved" in text
    assert "Second line." in text
    assert "RAW analysis scope" not in text
    # The MSAPO template itself is still the container.
    assert "MSAPO" in text
    assert "Labor and materials" in text
    assert "After-hours work" in text


def test_document_uses_reviewed_facility_address_and_vendor_overrides():
    path = generate_docx(
        _analysis(),
        [],
        [],
        facility_display="Unity Hospital",
        facility_address_display="1555 Long Pond Rd, Rochester, NY 14626",
        vendor_display="Reviewed Vendor LLC",
    )
    try:
        text = "\n".join(p.text for p in Document(str(path)).paragraphs)
    finally:
        path.unlink(missing_ok=True)

    assert "Facility: Unity Hospital" in text
    assert "1555 Long Pond Rd, Rochester, NY 14626" in text
    assert "Vendor: Reviewed Vendor LLC" in text
    assert "Crosby-Brownlie" not in text


def test_renderer_failure_is_reported_rather_than_returning_junk(monkeypatch):
    """A renderer that 'succeeds' without producing a PDF must not pass through.

    The bytes go straight into the Smartsheet attachment set, where a non-PDF
    payload would fail validation far from its cause.
    """
    fake = OUTPUT_DIR / "_not_really_a_pdf.pdf"
    fake.write_bytes(b"this is not a pdf")
    monkeypatch.setattr(
        "app.pdf_converter.convert_to_pdf", lambda _path: fake
    )

    with pytest.raises(PDFConversionError, match="valid PDF"):
        build_msapo_pdf(
            analysis=_analysis(),
            scope="Reviewed scope",
            inclusions=[],
            exclusions=[],
        )
    fake.unlink(missing_ok=True)


@pytest.mark.skipif(
    not _renderer_available(),
    reason="libreoffice-writer is unavailable in this environment",
)
def test_end_to_end_render_produces_a_pdf_and_leaves_no_intermediates():
    before = {p.name for p in Path(OUTPUT_DIR).iterdir()}

    payload = build_msapo_pdf(
        analysis=_analysis(),
        scope="REVIEWED scope the operator approved.",
        inclusions=["Labor and materials"],
        exclusions=["After-hours work"],
        facility_display="Rochester General Hospital",
    )

    assert payload.startswith(b"%PDF-")
    assert len(payload) > 1000

    import fitz

    with fitz.open(stream=payload, filetype="pdf") as document:
        text = "\n".join(page.get_text() for page in document)
    assert "MSAPO" in text
    assert "REVIEWED scope the operator approved" in text
    assert "RAW analysis scope" not in text

    after = {p.name for p in Path(OUTPUT_DIR).iterdir()}
    assert after == before, f"left intermediates behind: {sorted(after - before)}"


def test_each_conversion_gets_an_isolated_libreoffice_profile(monkeypatch):
    """Concurrency and crash-recovery both depend on this argument.

    Sharing the default profile under $HOME fails in production in ways a
    single-threaded test run never shows: two concurrent conversions contend for
    the same profile and the second never converts, and a run killed mid-flight
    leaves a lock that poisons the profile for every later run -- so the feature
    works once and then fails for the life of the container. The Calc path
    already isolates; this pins the Writer path to the same rule.
    """
    import subprocess

    from app import pdf_converter

    captured: dict[str, list[str]] = {}

    class _Result:
        returncode = 0
        stderr = ""
        stdout = ""

    def _fake_run(cmd, **kwargs):
        captured["cmd"] = list(cmd)
        # Produce the file the converter expects so it does not raise.
        (OUTPUT_DIR / f"{Path(cmd[-1]).stem}.pdf").write_bytes(b"%PDF-1.7\nstub")
        return _Result()

    monkeypatch.setattr(pdf_converter, "PDF_BACKEND", "libreoffice")
    monkeypatch.setattr(pdf_converter.shutil, "which", lambda _n: "/usr/bin/soffice")
    monkeypatch.setattr(subprocess, "run", _fake_run)

    source = OUTPUT_DIR / "_profile_probe.docx"
    document = Document()
    document.add_paragraph("probe")
    document.save(str(source))
    try:
        produced = pdf_converter.convert_to_pdf(source)
    finally:
        source.unlink(missing_ok=True)
    produced.unlink(missing_ok=True)

    profile_args = [a for a in captured["cmd"] if a.startswith("-env:UserInstallation=")]
    assert profile_args, (
        "no -env:UserInstallation argument: conversions would share the default "
        f"profile. command was: {captured['cmd']}"
    )
    assert profile_args[0].startswith("-env:UserInstallation=file://")
    assert "--headless" in captured["cmd"]
